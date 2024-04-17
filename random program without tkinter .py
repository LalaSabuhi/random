import docx
import os
import random
doc=docx.Document('C:\\Users\\e1-572g\\Downloads\\sual.docx')
yazilasi=docx.Document()
doc.paragraphs
#------------------------------------------------------------
a = [0] * 500
a1=int(input('1-ci sualin diapazonunun evveli:\t'))
a2=int(input('1-ci sualin diapazonunun sonu:\t'))
a3=int(input('2-ci sualin diapazonunun evveli:\t'))
a4=int(input('2-ci sualin diapazonunun sonu:\t'))
a5=int(input('3-cu sualin diapazonunun evveli:\t'))
a6=int(input('3-cu sualin diapazonunun sonu:\t'))
a7=int(input('4-cu sualin diapazonunun evveli:\t'))
a8=int(input('4-cu sualin diapazonunun sonu:\t'))
a9=int(input('5-ci sualin diapazonunun evveli:\t'))
a10=int(input('5-ci sualin diapazonunun sonu:\t'))
a11=int(input('yaradilacaq biletlerin sayi:'))
#------------------------------------------------------------
for i in range (1,a11+1):
    while True:
        n=random.randint(a1,a2)
        if(a[n]==0):
            break
        a[n]=1
#-------------------------------------------------------------
    p=1
    for u in range (a1,a2+1):
        p=p*a[u]
    if p==1:
        for v in range (a1,a2+1):
            a[v]=0
#--------------------------------------------------------------
    while True:
        k=random.randint(a3,a4)
        if(a[k]==0):
            break
        a[k]=1
#-------------------------------------------------------------
    p=1
    for u in range (a3,a4+1):
        p=p*a[u]
    if p==1:
        for v in range (a3,a4+1):
           a[v]=0
#--------------------------------------------------------------
    while True:
        m=random.randint(a5,a6)
        if(a[m]==0):
            break
        a[m]=1
#-------------------------------------------------------------
    p=1
    for u in range (a5,a6+1):
        p=p*a[u]
    if p==1:
        for v in range (a5,a6+1):
           a[v]=0
#--------------------------------------------------------------
    while True:
        s=random.randint(a7,a8)
        if (a[s]==0):
            break
        a[s]=1
#-------------------------------------------------------------
    p=1
    for u in range (a7,a8+1):
        p=p*a[u]
    if p==1:
        for v in range (a7,a8+1):
           a[v]=0
#--------------------------------------------------------------
    while True:
        t=random.randint(a9,a10)
        if  (a[t]==0):
            break
        a[t]=1
#-------------------------------------------------------------
    p=1
    for u in range (a9,a10+1):
        p=p*a[u]
    if p==1:
        for v in range (a9,a10+1):
           a[v]=0
#--------------------------------------------------------------------
    yazilasi.add_paragraph('  \n'+'\t\t\t\t Bilet '+str(i)+'  \n')
    yazilasi.add_paragraph('1.'+doc.paragraphs[n-1].text)
    yazilasi.add_paragraph("2."+doc.paragraphs[k-1].text)
    yazilasi.add_paragraph("3."+doc.paragraphs[m-1].text)
    yazilasi.add_paragraph("4."+doc.paragraphs[s-1].text)
    yazilasi.add_paragraph("5."+doc.paragraphs[t-1].text)
    yazilasi.save('C:\\Users\\e1-572g\\Desktop\\sual1.docx')
os.system("start C:\\Users\\e1-572g\\Desktop\\sual1.docx")
