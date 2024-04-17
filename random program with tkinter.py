Efrom tkinter import *
import docx
import os
import random
root=Tk()
root.title("Bilet Proqramı")
metn=Text(root,width=50, height=2,bg='gray')
metn.grid(row=20,column=0,columnspan=10)
def change(event):
    s=str(s1.get())
    doc=docx.Document(s)
    yazilasi=docx.Document()
    doc.paragraphs
    c=str(s2.get())
    a1=int(sual1.get())
    a2=int(sual2.get())
    a3=int(sual3.get())
    a4=int(sual4.get())
    a5=int(sual5.get())
    a6=int(sual6.get())
    a7=int(sual7.get())
    a8=int(sual8.get())
    a9=int(sual9.get())
    a10=int(sual10.get())
    a11=int(sual11.get())
    a = [0] * 500
    for i in range(1,a11+1):
        while True:
            n=random.randint(a1,a2)
            if(a[n]==0):
                break
            a[n]=1
    #-------------------------------------------------------------
        p=1
        for u in range (a1,a2+1):
            p=p*a[u]
        if p==1:
            for v in range (a1,a2+1):
                a[v]=0
    #--------------------------------------------------------------
        while True:
            k=random.randint(a3,a4)
            if(a[k]==0):
                break
            a[k]=1
    #-------------------------------------------------------------
        p=1
        for u in range (a3,a4+1):
            p=p*a[u]
        if p==1:
            for v in range (a3,a4+1):
                a[v]=0
    #--------------------------------------------------------------
        while True:
            m=random.randint(a5,a6)
            if(a[m]==0):
                break
            a[m]=1
    #-------------------------------------------------------------
        p=1
        for u in range (a5,a6+1):
            p=p*a[u]
        if p==1:
            for v in range (a5,a6+1):
                a[v]=0
    #4--------------------------------------------------------------
        while True:
            s=random.randint(a7,a8)
            if (a[s]==0):
                break
            a[s]=1
    #-------------------------------------------------------------
        p=1
        for u in range (a7,a8+1):
            p=p*a[u]
        if p==1:
            for v in range (a7,a8+1):
                a[v]=0
    #5--------------------------------------------------------------
        while True:
            t=random.randint(a9,a10)
            if (a[t]==0):
                break
            a[t]=1
    #-------------------------------------------------------------
        p=1
        for u in range (a9,a10+1):
            p=p*a[u]
        if p==1:
            for v in range (a9,a10+1):
                a[v]=0
    #--------------------------------------------------------------------
        yazilasi.add_paragraph('  \n'+'\t\t\t\t Bilet '+str(i)+'  \n')
        yazilasi.add_paragraph('1.'+doc.paragraphs[n-1].text)
        yazilasi.add_paragraph("2."+doc.paragraphs[k-1].text)
        yazilasi.add_paragraph("3."+doc.paragraphs[m-1].text)
        yazilasi.add_paragraph("4."+doc.paragraphs[s-1].text)
        yazilasi.add_paragraph("5."+doc.paragraphs[t-1].text)
        yazilasi.save(c)
    os.system(c)
#------------------------------------------------------------
s1=StringVar()
txt15=Entry(root,width=50,textvariable=s1)
txt15.grid(row=13,column=3,stick='w')

s2=StringVar()
txt16=Entry(root,width=50,textvariable=s2)
txt16.grid(row=14,column=3,stick='w')

sual1=StringVar()
txt1=Entry(root,width=10,textvariable=sual1)
txt1.grid(row=1,column=1,stick='w')
sual2=StringVar()
txt2=Entry(root,width=10,textvariable=sual2)
txt2.grid(row=1,column=2,stick='w')
sual3=StringVar()
txt3=Entry(root,width=10,textvariable=sual3)
txt3.grid(row=2,column=1,stick='w')
sual4=StringVar()
txt4=Entry(root,width=10,textvariable=sual4)
txt4.grid(row=2,column=2,stick='w')
sual5=StringVar()
txt5=Entry(root,width=10,textvariable=sual5)
txt5.grid(row=3,column=1,stick='w')
sual6=StringVar()
txt6=Entry(root,width=10,textvariable=sual6)
txt6.grid(row=3,column=2,stick='w')
sual7=StringVar()
txt7=Entry(root,width=10,textvariable=sual7)
txt7.grid(row=4,column=1,stick='w')
sual8=StringVar()
txt8=Entry(root,width=10,textvariable=sual8)
txt8.grid(row=4,column=2,stick='w')
sual9=StringVar()
txt9=Entry(root,width=10,textvariable=sual9)
txt9.grid(row=5,column=1,stick='w')
sual10=StringVar()
txt10=Entry(root,width=10,textvariable=sual10)
txt10.grid(row=5,column=2,stick='w')
sual11=StringVar()
txt11=Entry(root,width=10,textvariable=sual11)
txt11.grid(row=8,column=3,stick='w')
#------------------------------------------------------------
lbl_sual1=Label(root,text="sual1=",bg='yellow')
lbl_sual1.grid(row=1,column=0,stick='e')
lbl_sual2=Label(root,text="sual2=",bg='yellow')
lbl_sual2.grid(row=2,column=0,stick='e')
lbl_sual3=Label(root,text="sual3=",bg='yellow')
lbl_sual3.grid(row=3,column=0,stick='e')
lbl_sual4=Label(root,text="sual4=",bg='yellow')
lbl_sual4.grid(row=4,column=0,stick='e')
lbl_sual5=Label(root,text="sual5=",bg='yellow')
lbl_sual5.grid(row=5,column=0,stick='e')
lbl_sual11=Label(root,text="Biletlərin sayı=",bg='red')
lbl_sual11.grid(row=8,column=2,stick='e')
lbl_s1=Label(root,text="Sual_location=",bg='red')
lbl_s1.grid(row=13,column=2,stick='e')
lbl_s2=Label(root,text="Bilet_location=",bg='red')
lbl_s2.grid(row=14,column=2,stick='e')
btn=Button(root,text="Başla",bg='green')
btn.grid(row=15,column=0,stick='we',columnspan=50)
btn.bind('<Button-1>',change)
root.mainloop()
